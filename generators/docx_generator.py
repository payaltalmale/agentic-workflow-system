"""
DOCX Generator

Creates the final Word document from the approved draft.
Markdown-style headings (#, ##, ###) are converted into Word headings.
Documents are stored in data/generated_docs with a readable filename.
"""

from pathlib import Path
import re

from docx import Document
from docx.shared import Pt

OUTPUT_DIR = Path("data/generated_docs")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def _safe_filename(text: str) -> str:
    """
    Convert a title into a safe filename.
    """
    text = text.strip().replace(" ", "_")
    text = re.sub(r"[^\w\-]", "", text)
    return text[:60]  # limit filename length


def generate_docx(run_id: str, title: str, draft: str) -> str:
    doc = Document()

    doc.add_heading(title, level=0)

    for line in draft.split("\n"):
        stripped = line.strip()

        if not stripped:
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)

        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)

        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)

        else:
            p = doc.add_paragraph(stripped)
            p.style.font.size = Pt(11)

    safe_title = _safe_filename(title)

    filename = f"{safe_title}_{run_id}.docx"

    out_path = OUTPUT_DIR / filename

    doc.save(out_path)

    return str(out_path)